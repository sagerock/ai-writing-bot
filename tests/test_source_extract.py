import io
import unittest

from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from source_extract import extract


def _pdf_bytes(*page_texts):
    writer = PdfWriter()
    for page_text in page_texts:
        page = writer.add_blank_page(width=612, height=792)
        font = DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
            }
        )
        font_ref = writer._add_object(font)
        page[NameObject("/Resources")] = DictionaryObject(
            {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})}
        )
        stream = DecodedStreamObject()
        escaped = (
            page_text.replace("\\", "\\\\")
            .replace("(", "\\(")
            .replace(")", "\\)")
        )
        stream.set_data(f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode("ascii"))
        page[NameObject("/Contents")] = writer._add_object(stream)

    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def _docx_bytes():
    document = Document()
    document.add_paragraph("Opening paragraph")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Limit"
    table.cell(1, 1).text = "Three years"
    document.add_paragraph("Closing paragraph")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


class SourceExtractTests(unittest.TestCase):
    def assert_offsets_match(self, result):
        for entry in result["pages"]:
            self.assertTrue(result["text"][entry["start"] : entry["end"]].strip())

    def test_pdf_preserves_page_boundaries(self):
        result = extract(_pdf_bytes("First page", "Second page"), "application/pdf")

        self.assertEqual(result["kind"], "page")
        self.assertEqual([entry["page"] for entry in result["pages"]], [1, 2])
        self.assertEqual(
            [
                result["text"][entry["start"] : entry["end"]].strip()
                for entry in result["pages"]
            ],
            ["First page", "Second page"],
        )
        self.assert_offsets_match(result)

    def test_txt_maps_blank_line_separated_paragraphs(self):
        result = extract(
            b"First paragraph.\r\n\r\nSecond paragraph.\n",
            "text/plain; charset=utf-8",
        )

        self.assertEqual(result["kind"], "paragraph")
        self.assertEqual([entry["paragraph"] for entry in result["pages"]], [1, 2])
        self.assertEqual(result["text"], "First paragraph.\n\nSecond paragraph.\n")
        self.assert_offsets_match(result)

    def test_markdown_uses_paragraph_blocks(self):
        result = extract(b"# Heading\n\nBody text.\n", "text/markdown")

        self.assertEqual(len(result["pages"]), 2)
        self.assertEqual(
            result["text"][
                result["pages"][0]["start"] : result["pages"][0]["end"]
            ],
            "# Heading",
        )
        self.assert_offsets_match(result)

    def test_docx_preserves_paragraph_and_table_order(self):
        result = extract(
            _docx_bytes(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        self.assertEqual(result["kind"], "paragraph")
        self.assertLess(
            result["text"].index("Opening paragraph"),
            result["text"].index("Name\tValue"),
        )
        self.assertLess(
            result["text"].index("Name\tValue"),
            result["text"].index("Closing paragraph"),
        )
        self.assertEqual(len(result["pages"]), 3)
        self.assert_offsets_match(result)

    def test_csv_becomes_a_mapped_markdown_table(self):
        result = extract(b"name,note\nSmith,one|two\nJones,three\n", "text/csv")

        self.assertEqual(result["kind"], "paragraph")
        self.assertIn("| name | note |\n| --- | --- |", result["text"])
        self.assertIn(r"one\|two", result["text"])
        self.assertEqual(len(result["pages"]), 3)
        self.assert_offsets_match(result)

    def test_rejects_unknown_binary_types(self):
        with self.assertRaisesRegex(ValueError, "Unsupported content type"):
            extract(b"binary", "application/octet-stream")


if __name__ == "__main__":
    unittest.main()
